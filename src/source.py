import fitz  # Using PyMuPDF as it's faster and would be better if there are thousand+ resumes to be ranked
from pdf2image import convert_from_path
from docx import Document  # Using python-docx because it extracts with proper formatting and maintains text order
import pdfplumber
from pathlib import Path
import re
import nltk
from nltk.corpus import stopwords
from sentence_transformers import SentenceTransformer
import os

# Check if stopwords are already downloaded
try:
    nltk.data.find("corpora/stopwords")
except LookupError:
    nltk.download("stopwords")

stop_words = set(stopwords.words("english"))

def extract_text_from_resume(pdf_path):
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc]).strip()
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join([para.text.strip() for para in doc.paragraphs]).strip()
    return text

def extract_text(file_obj, filename):
    """Extracts text from a file object (PDF or DOCX) stored in memory."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        try:
            with pdfplumber.open(file_obj) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages]).strip()
            return text
        except Exception:
            return extract_text_from_image_pdf(file_obj)  # Use OCR fallback

    elif ext == ".docx":
        file_obj.seek(0)  # Reset pointer before reading
        doc = Document(file_obj)
        text = "\n".join([para.text.strip() for para in doc.paragraphs]).strip()
        return text

    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")

        
def preprocess_text(text, remove_stopwords=True):
    """Lowercases text, removes special characters, and optionally removes stopwords."""
    text = text.lower()
    text = re.sub(r"[^a-zA-Z0-9\s]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    
    if remove_stopwords:
        text = " ".join(word for word in text.split() if word not in stop_words)

    return text

# Load SBERT model
model = SentenceTransformer('all-mpnet-base-v2')

def rank_resumes(resume_files, job_description):
    """Ranks multiple resumes based on similarity to job description."""
    job_embedding = model.encode([job_description], normalize_embeddings=True)

    ranked_resumes = []
    
    for filename, file_obj in resume_files.items():
        try:
            resume_text = extract_text(file_obj, filename)
        except Exception as e:
            print(f"Error processing {filename}: {e}")
            continue  # Skip this resume
            
        resume_text = preprocess_text(resume_text)
        
        resume_embedding = model.encode([resume_text], normalize_embeddings=True)
        similarity_score = float(resume_embedding @ job_embedding.T)  # Using dot product 
        
        ranked_resumes.append((filename, similarity_score))

    # Sort resumes by similarity score in descending order
    ranked_resumes.sort(key=lambda x: x[1], reverse=True)
    print(ranked_resumes)  # Debugging step

    return ranked_resumes  # Return ranked results
